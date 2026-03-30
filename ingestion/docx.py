"""
ingestion/docx.py
=================
Extrator para documentos Word (.docx, .doc).
Usa python-docx. Preserva parágrafos e tabelas.
"""

from pathlib import Path
from typing import Any, Dict, List, Tuple

from .base import AbstractExtractor, ExtractionResult


class DocxExtractor(AbstractExtractor):
    supported_extensions = [".docx", ".doc"]

    def extract(self, file_path: Path) -> ExtractionResult:
        if not file_path.exists():
            return self._build_error_result(file_path, "docx", f"Arquivo não encontrado: {file_path}")

        try:
            from docx import Document
        except ImportError:
            return self._build_error_result(file_path, "docx", "python-docx não instalado (pip install python-docx)")

        try:
            text, pages, meta = self._extract(file_path)
            return ExtractionResult(
                text=self._clean_text(text),
                source_path=str(file_path),
                format="docx",
                page_count=meta.get("page_count", 1),
                extraction_method="python-docx",
                quality_score=1.0 if len(text) > 100 else 0.5,
                pages=pages,
                metadata=meta,
            )
        except Exception as exc:
            return self._build_error_result(file_path, "docx", str(exc))

    def _extract(self, path: Path) -> Tuple[str, List[Dict], Dict[str, Any]]:
        from docx import Document

        doc = Document(str(path))
        paragraphs: List[str] = []

        # Parágrafos de texto
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Conteúdo de tabelas (laudos, evoluções em grade)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n".join(paragraphs)
        pages = [{"page": 1, "text": full_text}]
        meta = {
            "page_count": 1,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        return full_text, pages, meta
